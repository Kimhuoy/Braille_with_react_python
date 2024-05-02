
from KhmerToBrailleLogic import translateKhmerToBraille
from BrailleToKhmerLogic import translateBrailleToKhmer
from KhmerToBrailleLogic import createUnicode
from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
from docx import Document
import requests, os
import fitz
from flask_cors import CORS
import logging

app = Flask(__name__)
CORS(app)
CORS(app, resources={r"/braille": {"origins": "http://127.0.0.1:5173"}})

# handle for upload text file to translate
@app.route('/updateTextFile', methods=['POST'])
def updateTextFile():   
    global uploaded_text
    if request.method == 'POST':
        try:
            if 'file' not in request.files:
                return jsonify({"error": "No file part"}), 400

            upload = request.files['file']
            if upload.filename == '':
                return jsonify({"error": "No selected file"}), 400

            if upload:
                file_path = os.path.join("uploads", secure_filename(upload.filename))
                upload.save(file_path)
                print(f"File saved at: {file_path}")
            file_extension = os.path.splitext(file_path)[1]

            if(file_extension == ".txt"):
                try:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        combined_lines = ""
                        for line in file:
                            combined_lines += line.strip() + "\n"
                            uploaded_text = combined_lines
                            print("text:", uploaded_text)
                except FileNotFoundError:
                    return 'File not found', 404
                except Exception as e:
                    return f'An error occurred: {str(e)}', 500
                
            if(file_extension ==".docx" or file_extension ==".doc"):
                if file_path:
                    try:
                        document = Document(file_path)
                        uploaded_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
                        print("docx: ",uploaded_text)
                    except Exception as e:
                        error_message = f"Error reading Word file: {e}"

            if file_extension == ".pdf":
                try:
                    with fitz.open(file_path) as pdf_document:
                        combined_text = ""
                        for page_num in range(pdf_document.page_count):
                            page = pdf_document[page_num]
                            text = page.get_text("text", )
                            combined_text += str(text)
                        
                        uploaded_text = combined_text
                        print("pdf: ", uploaded_text)
                except Exception as e:
                    error_message = f"Error reading PDF file: {e}"
                    return jsonify({"error pdf": error_message})
            
        except Exception as e:
            logging.error(f"Error processing file: {e}")
            return jsonify({"error": "Internal Server Error"}), 500
        
    text = translateKhmerToBraille(uploaded_text)
    print(text)
    return jsonify({"result": text, "uploaded": uploaded_text})

#handle to upload text to translate to braille
@app.route('/text', methods=['POST'])
def handleTextToBraille():
    data = request.get_json()
    uploaded_text = data.get("text", "")
    print("Upload text: ",uploaded_text)
    text = translateKhmerToBraille(uploaded_text)
    print(text)
    unicode = createUnicode();
    return jsonify({"result":text, "unicodeData":unicode})

#handle to upload braille file to translate 
@app.route('/updatebraille', methods=['POST'])
def updateBrailleFile():   
    global uploaded_braille
    if request.method == 'POST':
        try:
            if 'file' not in request.files:
                return jsonify({"error": "No file part"}), 400

            upload = request.files['file']
            if upload.filename == '':
                return jsonify({"error": "No selected file"}), 400

            if upload:
                file_path = os.path.join("uploads", secure_filename(upload.filename))
                upload.save(file_path)
                print(f"File saved at: {file_path}")
            file_extension = os.path.splitext(file_path)[1]

            if(file_extension == ".txt"):
                try:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        combined_lines = ""
                        for line in file:
                            combined_lines += line.strip()
                            uploaded_braille = combined_lines
                            print("text:", uploaded_braille)
                except FileNotFoundError:
                    return 'File not found', 404
                except Exception as e:
                    return f'An error occurred: {str(e)}', 500
                
            if(file_extension ==".docx" or file_extension ==".doc"):
                if file_path:
                    try:
                        document = Document(file_path)
                        uploaded_braille = "\n".join(paragraph.text for paragraph in document.paragraphs)
                        print("docx: ",uploaded_braille)
                    except Exception as e:
                        error_message = f"Error reading Word file: {e}"

            if file_extension == ".pdf":
                try:
                    with fitz.open(file_path) as pdf_document:
                        combined_text = ""
                        for page_num in range(pdf_document.page_count):
                            page = pdf_document[page_num]
                            text = page.get_text("text", )
                            combined_text += str(text)
                        
                        uploaded_braille = combined_text
                        print("pdf: ", uploaded_braille)
                except Exception as e:
                    error_message = f"Error reading PDF file: {e}"
                    return jsonify({"error pdf": error_message})
            
        except Exception as e:
            logging.error(f"Error processing file: {e}")
            return jsonify({"error": "Internal Server Error"}), 500
        
    braille = translateBrailleToKhmer(uploaded_braille)
    print(braille)
    return jsonify({"result": braille, "uploaded": uploaded_braille})

@app.route('/braille', methods=['POST'])
def handleBrailleToText():
    try:
        data = request.get_json()
        if data is None or "braille" not in data:
            return jsonify({"error": "Invalid request data"}), 400
        uploaded_braille = data.get("braille", "")
        logging.info("Received braille: %s", uploaded_braille)
        try:
            braille = translateBrailleToKhmer(uploaded_braille)
        except Exception as e:  
            return jsonify({"error": str(e)}), 500
        print("braille:", braille)
        return jsonify({"result": braille})
    except Exception as e:
        logging.exception("Error processing request")
        return jsonify({"error": "Internal Server Error"}), 500

if __name__ == '__main__':
    app.run(debug=True)
